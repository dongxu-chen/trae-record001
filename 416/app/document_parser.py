import os
import re
from typing import List, Optional, Dict, Tuple
from dataclasses import dataclass, field
from langchain_core.documents import Document


@dataclass
class TextSegment:
    content: str
    page: Optional[int] = None
    start_line: int = 0
    end_line: int = 0
    start_char: int = 0
    end_char: int = 0
    content_type: str = "text"


@dataclass
class TableData:
    table_id: str
    headers: List[str]
    rows: List[List[str]]
    page: Optional[int] = None
    start_line: int = 0
    end_line: int = 0


@dataclass
class ChartData:
    chart_id: str
    chart_type: str
    title: str
    data_points: List[Dict]
    page: Optional[int] = None
    start_line: int = 0
    end_line: int = 0


@dataclass
class ParsedDocument:
    content: str
    metadata: dict
    filename: str
    file_type: str
    segments: List[TextSegment] = field(default_factory=list)
    tables: List[TableData] = field(default_factory=list)
    charts: List[ChartData] = field(default_factory=list)


class DocumentParser:
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "text",
    }

    CHART_KEYWORDS = [
        "图", "表", "chart", "graph", "figure", "fig.",
        "趋势", "分布", "统计", "对比", "占比", "增长"
    ]

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        ext = os.path.splitext(filename)[1].lower()
        return ext in cls.SUPPORTED_EXTENSIONS

    @classmethod
    def get_file_type(cls, filename: str) -> str:
        ext = os.path.splitext(filename)[1].lower()
        return cls.SUPPORTED_EXTENSIONS.get(ext, "unknown")

    @classmethod
    def parse(cls, file_path: str, filename: str) -> ParsedDocument:
        file_type = cls.get_file_type(filename)
        
        if file_type == "pdf":
            content, metadata, segments, tables, charts = cls._parse_pdf(file_path)
        elif file_type == "docx":
            content, metadata, segments, tables, charts = cls._parse_docx(file_path)
        elif file_type == "markdown":
            content, metadata, segments, tables, charts = cls._parse_markdown(file_path)
        elif file_type == "text":
            content, metadata, segments, tables, charts = cls._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        metadata["filename"] = filename
        metadata["file_type"] = file_type
        metadata["file_path"] = file_path
        metadata["table_count"] = len(tables)
        metadata["chart_count"] = len(charts)

        return ParsedDocument(
            content=content,
            metadata=metadata,
            filename=filename,
            file_type=file_type,
            segments=segments,
            tables=tables,
            charts=charts,
        )

    @classmethod
    def _parse_pdf(
        cls, file_path: str
    ) -> Tuple[str, dict, List[TextSegment], List[TableData], List[ChartData]]:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        content_parts = []
        segments = []
        tables = []
        charts = []
        current_char_pos = 0
        current_line_pos = 0
        table_idx = 0
        chart_idx = 0

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            lines = text.split("\n")
            
            page_content = f"--- Page {page_num} ---\n{text}\n"
            content_parts.append(page_content)
            
            for line_idx, line in enumerate(lines):
                if line.strip():
                    content_type = "text"
                    
                    if cls._is_table_line(line):
                        table_idx += 1
                        content_type = "table"
                    elif cls._is_chart_reference(line):
                        chart_idx += 1
                        content_type = "chart"
                    
                    line_end = current_char_pos + len(line) + 1
                    segment = TextSegment(
                        content=line,
                        page=page_num,
                        start_line=current_line_pos + line_idx,
                        end_line=current_line_pos + line_idx,
                        start_char=current_char_pos,
                        end_char=line_end,
                        content_type=content_type,
                    )
                    segments.append(segment)
                    
                    if content_type == "table":
                        table = cls._extract_table_from_pdf_line(
                            line, page_num, current_line_pos + line_idx, table_idx
                        )
                        if table:
                            tables.append(table)
                    elif content_type == "chart":
                        chart = cls._extract_chart_from_line(
                            line, page_num, current_line_pos + line_idx, chart_idx
                        )
                        if chart:
                            charts.append(chart)
                
                current_char_pos += len(line) + 1
            
            current_line_pos += len(lines) + 2
            current_char_pos += len(f"--- Page {page_num} ---\n") + 1

        content = "\n".join(content_parts)
        metadata = {
            "page_count": len(reader.pages),
            "total_lines": current_line_pos,
            "total_chars": len(content),
        }
        return content, metadata, segments, tables, charts

    @classmethod
    def _parse_docx(
        cls, file_path: str
    ) -> Tuple[str, dict, List[TextSegment], List[TableData], List[ChartData]]:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        content_parts = []
        segments = []
        tables = []
        charts = []
        current_char_pos = 0
        current_line_pos = 0

        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                lines = para.text.split("\n")
                for line in lines:
                    if line.strip():
                        content_type = "text"
                        if cls._is_chart_reference(line):
                            content_type = "chart"
                        
                        line_end = current_char_pos + len(line)
                        segments.append(TextSegment(
                            content=line,
                            start_line=current_line_pos,
                            end_line=current_line_pos,
                            start_char=current_char_pos,
                            end_char=line_end,
                            content_type=content_type,
                        ))
                        
                        if content_type == "chart":
                            chart = cls._extract_chart_from_line(
                                line, None, current_line_pos, len(charts) + 1
                            )
                            if chart:
                                charts.append(chart)
                    
                    content_parts.append(line)
                    current_char_pos += len(line) + 1
                    current_line_pos += 1
            else:
                content_parts.append("")
                current_char_pos += 1
                current_line_pos += 1

        for table_idx, table in enumerate(doc.tables):
            table_data = cls._parse_docx_table(
                table, table_idx + 1, current_line_pos
            )
            if table_data:
                tables.append(table_data)
                
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        line_end = current_char_pos + len(row_text)
                        segments.append(TextSegment(
                            content=row_text,
                            start_line=current_line_pos,
                            end_line=current_line_pos,
                            start_char=current_char_pos,
                            end_char=line_end,
                            content_type="table",
                        ))
                        content_parts.append(row_text)
                        current_char_pos += len(row_text) + 1
                        current_line_pos += 1
                
                table_data.start_line = current_line_pos - len(table.rows)
                table_data.end_line = current_line_pos - 1

        content = "\n".join(content_parts)
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "total_lines": current_line_pos,
            "total_chars": len(content),
        }
        return content, metadata, segments, tables, charts

    @classmethod
    def _parse_docx_table(
        cls, table, table_idx: int, start_line: int
    ) -> Optional[TableData]:
        if not table.rows:
            return None
        
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        
        return TableData(
            table_id=f"table_{table_idx}",
            headers=headers,
            rows=rows,
            start_line=start_line,
            end_line=start_line + len(table.rows) - 1,
        )

    @classmethod
    def _parse_markdown(
        cls, file_path: str
    ) -> Tuple[str, dict, List[TextSegment], List[TableData], List[ChartData]]:
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        
        return cls._process_lines_with_tables(lines, file_path)

    @classmethod
    def _parse_text(
        cls, file_path: str
    ) -> Tuple[str, dict, List[TextSegment], List[TableData], List[ChartData]]:
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        
        return cls._process_lines_with_tables(lines, file_path)

    @classmethod
    def _process_lines_with_tables(
        cls, lines: List[str], file_path: str
    ) -> Tuple[str, dict, List[TextSegment], List[TableData], List[ChartData]]:
        content_parts = []
        segments = []
        tables = []
        charts = []
        current_char_pos = 0
        table_idx = 0
        chart_idx = 0

        for line_idx, line in enumerate(lines):
            clean_line = line.rstrip("\n")
            if clean_line.strip():
                content_type = "text"
                
                if cls._is_markdown_table(clean_line):
                    table_idx += 1
                    content_type = "table"
                    table = cls._parse_markdown_table(
                        clean_line, line_idx, table_idx
                    )
                    if table:
                        tables.append(table)
                elif cls._is_chart_reference(clean_line):
                    chart_idx += 1
                    content_type = "chart"
                    chart = cls._extract_chart_from_line(
                        clean_line, None, line_idx, chart_idx
                    )
                    if chart:
                        charts.append(chart)
                
                line_end = current_char_pos + len(clean_line)
                segments.append(TextSegment(
                    content=clean_line,
                    start_line=line_idx,
                    end_line=line_idx,
                    start_char=current_char_pos,
                    end_char=line_end,
                    content_type=content_type,
                ))
            
            content_parts.append(clean_line)
            current_char_pos += len(clean_line) + 1

        content = "\n".join(content_parts)
        metadata = {
            "file_size": os.path.getsize(file_path),
            "total_lines": len(lines),
            "total_chars": len(content),
        }
        return content, metadata, segments, tables, charts

    @classmethod
    def _is_table_line(cls, text: str) -> bool:
        return bool(re.match(r'^\s*\|.*\|\s*$', text)) or \
               bool(re.match(r'^\s*[\-\+]+\s*$', text))

    @classmethod
    def _is_markdown_table(cls, text: str) -> bool:
        return bool(re.match(r'^\s*\|.*\|\s*$', text))

    @classmethod
    def _is_chart_reference(cls, text: str) -> bool:
        text_lower = text.lower()
        return any(kw in text_lower for kw in cls.CHART_KEYWORDS)

    @classmethod
    def _parse_markdown_table(
        cls, text: str, line_idx: int, table_idx: int
    ) -> Optional[TableData]:
        cells = re.split(r'\|', text.strip())
        cells = [c.strip() for c in cells if c.strip()]
        
        if len(cells) >= 2:
            return TableData(
                table_id=f"table_{table_idx}",
                headers=cells,
                rows=[],
                start_line=line_idx,
                end_line=line_idx,
            )
        return None

    @classmethod
    def _extract_table_from_pdf_line(
        cls, text: str, page: int, line: int, table_idx: int
    ) -> Optional[TableData]:
        cells = re.split(r'\s{2,}|\t', text.strip())
        cells = [c.strip() for c in cells if c.strip()]
        
        if len(cells) >= 2:
            return TableData(
                table_id=f"table_{table_idx}",
                headers=cells,
                rows=[],
                page=page,
                start_line=line,
                end_line=line,
            )
        return None

    @classmethod
    def _extract_chart_from_line(
        cls, text: str, page: Optional[int], line: int, chart_idx: int
    ) -> Optional[ChartData]:
        chart_types = ["折线图", "柱状图", "饼图", "散点图", "面积图", 
                      "line chart", "bar chart", "pie chart", "scatter", "area"]
        chart_type = "unknown"
        
        for ct in chart_types:
            if ct in text.lower():
                chart_type = ct
                break
        
        title_match = re.search(r'(?:图|图表|figure|fig\.?)\s*\d*[\.:：]\s*(.+?)(?:\n|$)', text, re.I)
        title = title_match.group(1) if title_match else text[:50]
        
        return ChartData(
            chart_id=f"chart_{chart_idx}",
            chart_type=chart_type,
            title=title,
            data_points=[],
            page=page,
            start_line=line,
            end_line=line,
        )

    @classmethod
    def find_segment_for_range(
        cls,
        segments: List[TextSegment],
        start_char: int,
        end_char: int,
    ) -> Tuple[Optional[int], Optional[int], Optional[int]]:
        page = None
        start_line = None
        end_line = None

        for seg in segments:
            if seg.start_char <= end_char and seg.end_char >= start_char:
                if page is None and seg.page is not None:
                    page = seg.page
                if start_line is None or seg.start_line < start_line:
                    start_line = seg.start_line
                if end_line is None or seg.end_line > end_line:
                    end_line = seg.end_line

        return page, start_line, end_line

    @classmethod
    def to_langchain_documents(cls, parsed_doc: ParsedDocument) -> List[Document]:
        metadata = parsed_doc.metadata.copy()
        
        if parsed_doc.tables:
            table_summaries = []
            for table in parsed_doc.tables:
                summary = f"表格 {table.table_id}: {len(table.headers)}列, {len(table.rows)}行"
                if table.headers:
                    summary += f" 列名: {', '.join(table.headers[:5])}"
                table_summaries.append(summary)
            metadata["table_summaries"] = "\n".join(table_summaries)
        
        if parsed_doc.charts:
            chart_summaries = []
            for chart in parsed_doc.charts:
                summary = f"图表 {chart.chart_id}: {chart.chart_type} - {chart.title}"
                chart_summaries.append(summary)
            metadata["chart_summaries"] = "\n".join(chart_summaries)
        
        doc = Document(
            page_content=parsed_doc.content,
            metadata=metadata,
        )
        return [doc]
